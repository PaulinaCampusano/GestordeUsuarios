"""
Genera el documento Word con el modelo de base de datos de GestionUsuario.
Uso: python3 generar_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Colores corporativos ────────────────────────────────────────────────────────
AZUL_OSCURO   = RGBColor(0x1E, 0x3A, 0x5F)   # encabezados principales
AZUL_MEDIO    = RGBColor(0x2E, 0x6D, 0xA4)   # encabezados de tabla
AZUL_CLARO    = RGBColor(0xDB, 0xE9, 0xF5)   # fila de encabezado de tabla
GRIS_FILA     = RGBColor(0xF5, 0xF5, 0xF5)   # filas alternas
VERDE         = RGBColor(0x1B, 0x7A, 0x34)   # PK / FK highlights
BLANCO        = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, color: RGBColor):
    """Aplica color de fondo a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    # RGBColor hereda de bytes — acceder por índice
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_border_to_table(table):
    """Añade bordes visibles a todas las celdas de la tabla."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "AAAAAA")
                tcBorders.append(border)
            tcPr.append(tcBorders)


def cell_text(cell, text, bold=False, color=None, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def make_table(doc, headers, rows, col_widths=None):
    """Crea una tabla estilizada."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Encabezado
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, AZUL_MEDIO)
        cell_text(cell, h, bold=True, color=BLANCO, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # Filas de datos
    for r_idx, row_data in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = GRIS_FILA if r_idx % 2 == 0 else BLANCO
        for c_idx, val in enumerate(row_data):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, bg)
            bold = c_idx == 0  # primera columna en negrita
            cell_text(cell, str(val), bold=bold, size=9)

    add_border_to_table(table)

    # Ancho de columnas
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[j])

    return table


# ══════════════════════════════════════════════════════════════════════════════
# CONSTRUCCIÓN DEL DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# ── Márgenes de página ────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Portada ───────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

titulo = doc.add_paragraph()
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titulo.add_run("MODELO DE BASE DE DATOS")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = AZUL_OSCURO

subtitulo = doc.add_paragraph()
subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitulo.add_run("GestionUsuario — Auth & User Management API")
run2.font.size = Pt(14)
run2.font.color.rgb = AZUL_MEDIO

doc.add_paragraph()
linea = doc.add_paragraph()
linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = linea.add_run("Spring Boot 3.3  ·  MySQL 8  ·  JPA / Hibernate")
run3.font.size = Pt(11)
run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

# ── Introducción ──────────────────────────────────────────────────────────────
doc.add_heading("1. Introducción", level=1)
intro = doc.add_paragraph(
    "El sistema GestionUsuario utiliza una base de datos relacional MySQL 8 llamada "
    "db_gestion_usuario. El esquema está gestionado por Hibernate con la política "
    "ddl-auto: update, lo que significa que las tablas se crean y actualizan "
    "automáticamente al arrancar la aplicación."
)
intro.style.font.size = Pt(11)

doc.add_paragraph(
    "El modelo contempla 6 tablas que cubren: autenticación de usuarios, control de "
    "acceso basado en roles (RBAC), una matriz de permisos granular por módulo y "
    "recuperación segura de contraseñas."
)

doc.add_page_break()

# ── Resumen de tablas ─────────────────────────────────────────────────────────
doc.add_heading("2. Resumen de Tablas", level=1)

resumen_headers = ["Tabla", "Entidad Java", "Propósito"]
resumen_rows = [
    ["USUARIO",               "User",               "Usuarios del sistema con credenciales y estado de cuenta"],
    ["ROL",                   "Role",               "Catálogo de roles (ROLE_ADMIN, ROLE_USER, ROLE_SUPERVISOR)"],
    ["MODULO",                "Modulo",             "Módulos del frontend que pueden protegerse con permisos"],
    ["ACCESO",                "Acceso",             "Tipos de operación: CREAR, LEER, ACTUALIZAR, ELIMINAR"],
    ["ROL_MODULO",            "RoleModulo",         "Tabla pivot — Matriz de permisos (rol + módulo + acceso)"],
    ["PASSWORD_RESET_TOKEN",  "PasswordResetToken", "Tokens UUID de un solo uso para recuperación de contraseña"],
]
make_table(doc, resumen_headers, resumen_rows, col_widths=[4.5, 4, 8.5])

doc.add_paragraph()

# ── Relaciones ────────────────────────────────────────────────────────────────
doc.add_heading("3. Relaciones entre Tablas", level=1)

rel_headers = ["Relación", "Cardinalidad", "Descripción"]
rel_rows = [
    ["ROL → USUARIO",              "1 : N",  "Un rol puede ser asignado a muchos usuarios. Cada usuario tiene exactamente 1 rol."],
    ["ROL → ROL_MODULO",           "1 : N",  "Un rol puede tener múltiples entradas en la matriz de permisos."],
    ["MODULO → ROL_MODULO",        "1 : N",  "Un módulo puede aparecer en múltiples entradas de permisos."],
    ["ACCESO → ROL_MODULO",        "1 : N",  "Un tipo de acceso puede asociarse a múltiples permisos."],
    ["USUARIO → PASSWORD_RESET_TOKEN", "1 : N", "Un usuario puede tener tokens de reset (máximo 1 activo a la vez)."],
]
make_table(doc, rel_headers, rel_rows, col_widths=[5, 3, 9])

doc.add_paragraph()
nota = doc.add_paragraph(
    "⚠  La clave primaria de ROL_MODULO es compuesta: (id_rol, id_modulo, id_acceso). "
    "Esto garantiza que no existan permisos duplicados para la misma combinación de rol, módulo y acción."
)
nota.runs[0].font.size = Pt(10)
nota.runs[0].font.italic = True

doc.add_page_break()

# ── Detalle de tablas ─────────────────────────────────────────────────────────
doc.add_heading("4. Detalle de Tablas", level=1)

# ── TABLA: USUARIO ─────────────────────────────────────────────────────────────
doc.add_heading("4.1  USUARIO", level=2)
doc.add_paragraph(
    "Tabla principal del sistema. Almacena las credenciales y datos personales de cada usuario. "
    "La contraseña se guarda como hash BCrypt (costo 10). El campo activo actúa como "
    "soft-delete — un usuario inactivo no puede autenticarse."
)

usuario_headers = ["Columna", "Tipo", "Restricción", "Descripción"]
usuario_rows = [
    ["id_usuario",        "BIGINT",   "PK, AUTO_INCREMENT",    "Identificador único del usuario"],
    ["id_rol",            "BIGINT",   "FK → ROL",              "Rol asignado al usuario"],
    ["nombre_completo",   "VARCHAR",  "NOT NULL",              "Nombre y apellidos completos"],
    ["email",             "VARCHAR",  "UNIQUE, NOT NULL",      "Correo electrónico / username de login"],
    ["password_hash",     "VARCHAR",  "NOT NULL",              "Contraseña encriptada con BCrypt"],
    ["rut",               "VARCHAR",  "UNIQUE, NULLABLE",      "RUT chileno (opcional)"],
    ["telefono",          "VARCHAR",  "NULLABLE",              "Número de teléfono de contacto"],
    ["activo",            "BOOLEAN",  "DEFAULT true",          "Estado de la cuenta (false = deshabilitado)"],
    ["token_recuperacion","VARCHAR",  "NULLABLE",              "Campo legacy (la tabla PASSWORD_RESET_TOKEN es la principal)"],
    ["expira_token",      "DATETIME", "NULLABLE",              "Expiración del token legacy"],
    ["fecha_creacion",    "DATETIME", "AUTO, NOT UPDATABLE",   "Fecha y hora de creación del registro"],
]
make_table(doc, usuario_headers, usuario_rows, col_widths=[4.5, 2.5, 4, 6])

doc.add_paragraph()

# ── TABLA: ROL ─────────────────────────────────────────────────────────────────
doc.add_heading("4.2  ROL", level=2)
doc.add_paragraph(
    "Catálogo de roles del sistema. Los nombres de rol siguen la convención de Spring Security "
    "(prefijo ROLE_). El sistema auto-genera el prefijo si no se incluye al crear un rol."
)

rol_headers = ["Columna", "Tipo", "Restricción", "Descripción"]
rol_rows = [
    ["id_rol",      "BIGINT",  "PK, AUTO_INCREMENT", "Identificador único del rol"],
    ["nombre",      "VARCHAR", "UNIQUE, NOT NULL",   "Nombre del rol (ej. ROLE_ADMIN, ROLE_USER, ROLE_SUPERVISOR)"],
    ["descripcion", "VARCHAR", "NULLABLE",           "Descripción del propósito del rol"],
]
make_table(doc, rol_headers, rol_rows, col_widths=[4.5, 2.5, 4, 6])

doc.add_paragraph()

seed_headers = ["nombre", "descripcion"]
seed_rows = [
    ["ROLE_ADMIN",      "Administrador del sistema"],
    ["ROLE_USER",       "Usuario estándar"],
    ["ROLE_SUPERVISOR", "Supervisor de acceso de solo lectura"],
]
p = doc.add_paragraph("Datos semilla iniciales:")
p.runs[0].bold = True
p.runs[0].font.size = Pt(10)
make_table(doc, seed_headers, seed_rows, col_widths=[5, 12])

doc.add_paragraph()

# ── TABLA: MODULO ──────────────────────────────────────────────────────────────
doc.add_heading("4.3  MODULO", level=2)
doc.add_paragraph(
    "Representa las secciones o pantallas del frontend que pueden ser controladas "
    "mediante permisos. El campo ruta_frontend es consumido directamente por el cliente "
    "para construir la navegación según los permisos del usuario autenticado."
)

mod_headers = ["Columna", "Tipo", "Restricción", "Descripción"]
mod_rows = [
    ["id_modulo",     "BIGINT",  "PK, AUTO_INCREMENT", "Identificador único del módulo"],
    ["nombre",        "VARCHAR", "NULLABLE",           "Nombre descriptivo del módulo"],
    ["descripcion",   "VARCHAR", "NULLABLE",           "Descripción del propósito del módulo"],
    ["ruta_frontend", "VARCHAR", "NULLABLE",           "Ruta del frontend asociada (ej. /usuarios, /roles)"],
]
make_table(doc, mod_headers, mod_rows, col_widths=[4.5, 2.5, 4, 6])

doc.add_paragraph()

p2 = doc.add_paragraph("Datos semilla iniciales:")
p2.runs[0].bold = True
p2.runs[0].font.size = Pt(10)
mod_seed_rows = [
    ["Gestión de Usuarios", "Módulo principal de administración de usuarios", "/usuarios"],
    ["Gestión de Roles",    "Módulo para administrar roles y permisos",       "/roles"],
]
make_table(doc, ["nombre", "descripcion", "ruta_frontend"], mod_seed_rows, col_widths=[5, 8, 4])

doc.add_paragraph()

# ── TABLA: ACCESO ──────────────────────────────────────────────────────────────
doc.add_heading("4.4  ACCESO", level=2)
doc.add_paragraph(
    "Define los tipos de operación que se pueden permitir sobre un módulo. "
    "Se sigue el modelo CRUD estándar."
)

acc_headers = ["Columna", "Tipo", "Restricción", "Descripción"]
acc_rows = [
    ["id_acceso",   "BIGINT",  "PK, AUTO_INCREMENT", "Identificador único del tipo de acceso"],
    ["nombre",      "VARCHAR", "NULLABLE",           "Nombre de la acción (CREAR, LEER, ACTUALIZAR, ELIMINAR)"],
    ["descripcion", "VARCHAR", "NULLABLE",           "Descripción de lo que permite esta acción"],
]
make_table(doc, acc_headers, acc_rows, col_widths=[4.5, 2.5, 4, 6])

doc.add_paragraph()

p3 = doc.add_paragraph("Datos semilla iniciales:")
p3.runs[0].bold = True
p3.runs[0].font.size = Pt(10)
acc_seed_rows = [
    ["CREAR",      "Permite crear registros"],
    ["LEER",       "Permite visualizar información"],
    ["ACTUALIZAR", "Permite modificar registros"],
    ["ELIMINAR",   "Permite eliminar o desactivar registros"],
]
make_table(doc, ["nombre", "descripcion"], acc_seed_rows, col_widths=[5, 12])

doc.add_paragraph()

doc.add_page_break()

# ── TABLA: ROL_MODULO ──────────────────────────────────────────────────────────
doc.add_heading("4.5  ROL_MODULO", level=2)
doc.add_paragraph(
    "Tabla pivot que implementa la matriz de permisos del sistema. "
    "Relaciona un ROL con un MODULO y un ACCESO, definiendo exactamente qué "
    "operaciones puede realizar cada rol sobre cada módulo del sistema."
)

rm_headers = ["Columna", "Tipo", "Restricción", "Descripción"]
rm_rows = [
    ["id_rol",    "BIGINT", "PK + FK → ROL",    "Rol al que aplica este permiso"],
    ["id_modulo", "BIGINT", "PK + FK → MODULO", "Módulo sobre el que aplica el permiso"],
    ["id_acceso", "BIGINT", "PK + FK → ACCESO", "Tipo de acceso (operación) permitida"],
]
make_table(doc, rm_headers, rm_rows, col_widths=[4.5, 2.5, 4, 6])

doc.add_paragraph()

p4 = doc.add_paragraph("Matriz de permisos semilla:")
p4.runs[0].bold = True
p4.runs[0].font.size = Pt(10)
perm_rows = [
    ["ROLE_ADMIN",      "Gestión de Usuarios", "CREAR"],
    ["ROLE_ADMIN",      "Gestión de Usuarios", "LEER"],
    ["ROLE_ADMIN",      "Gestión de Usuarios", "ACTUALIZAR"],
    ["ROLE_ADMIN",      "Gestión de Usuarios", "ELIMINAR"],
    ["ROLE_ADMIN",      "Gestión de Roles",    "CREAR"],
    ["ROLE_ADMIN",      "Gestión de Roles",    "LEER"],
    ["ROLE_ADMIN",      "Gestión de Roles",    "ACTUALIZAR"],
    ["ROLE_ADMIN",      "Gestión de Roles",    "ELIMINAR"],
    ["ROLE_SUPERVISOR", "Gestión de Usuarios", "LEER"],
    ["ROLE_SUPERVISOR", "Gestión de Roles",    "LEER"],
]
make_table(doc, ["ROL", "MÓDULO", "ACCESO"], perm_rows, col_widths=[5, 6, 6])

doc.add_paragraph()

# ── TABLA: PASSWORD_RESET_TOKEN ────────────────────────────────────────────────
doc.add_heading("4.6  PASSWORD_RESET_TOKEN", level=2)
doc.add_paragraph(
    "Almacena los tokens de recuperación de contraseña. Cada token es un UUID aleatorio "
    "con un tiempo de vida (TTL) de 15 minutos. Al solicitar un nuevo token, el anterior "
    "se elimina automáticamente, garantizando que solo exista 1 token activo por usuario. "
    "Una vez utilizado el token, también se elimina para prevenir su reutilización."
)

prt_headers = ["Columna", "Tipo", "Restricción", "Descripción"]
prt_rows = [
    ["id",               "BIGINT",   "PK, AUTO_INCREMENT", "Identificador único del token"],
    ["token",            "VARCHAR",  "UNIQUE, NOT NULL",   "UUID aleatorio generado al solicitar reset"],
    ["id_usuario",       "BIGINT",   "FK → USUARIO",       "Usuario propietario del token"],
    ["fecha_expiracion", "DATETIME", "NOT NULL",           "Fecha/hora límite de validez (creación + 15 min)"],
]
make_table(doc, prt_headers, prt_rows, col_widths=[4.5, 2.5, 4, 6])

doc.add_paragraph()

# ── Usuarios de prueba ────────────────────────────────────────────────────────
doc.add_page_break()
doc.add_heading("5. Usuarios de Prueba (DataSeeder)", level=1)
doc.add_paragraph(
    "Al arrancar la aplicación por primera vez, el componente DataSeeder crea "
    "automáticamente los siguientes usuarios con sus credenciales y roles asignados:"
)

ux_headers = ["Email", "Contraseña", "Rol", "Activo", "Descripción"]
ux_rows = [
    ["admin@gestion.com",      "AdminPass123!",   "ROLE_ADMIN",      "Sí", "Administrador con acceso total al sistema"],
    ["user@gestion.com",       "UserPass123!",    "ROLE_USER",       "Sí", "Usuario estándar sin permisos de módulo"],
    ["supervisor@gestion.com", "Supervisor123!",  "ROLE_SUPERVISOR", "Sí", "Supervisor con permisos de solo lectura"],
]
make_table(doc, ux_headers, ux_rows, col_widths=[5, 4, 4, 2, 7])

doc.add_paragraph()
nota2 = doc.add_paragraph(
    "ℹ  Los usuarios registrados mediante el endpoint público /api/auth/register quedan "
    "con activo = false y deben ser activados manualmente por un administrador "
    "usando el endpoint PATCH /api/admin/users/{id}/toggle-status."
)
nota2.runs[0].font.size = Pt(10)
nota2.runs[0].font.italic = True

# ── Política de contraseñas ───────────────────────────────────────────────────
doc.add_heading("6. Política de Contraseñas", level=1)
doc.add_paragraph(
    "Toda contraseña almacenada en el sistema debe cumplir las siguientes reglas, "
    "validadas por la clase PasswordPolicy antes de ser encriptada:"
)

pol_rows = [
    ["Longitud mínima",        "8 caracteres"],
    ["Letra mayúscula",        "Al menos 1 letra mayúscula (A-Z)"],
    ["Letra minúscula",        "Al menos 1 letra minúscula (a-z)"],
    ["Número",                 "Al menos 1 dígito numérico (0-9)"],
    ["Carácter especial",      "Al menos 1 carácter especial (!@#$%^&* etc.)"],
    ["Encriptación",           "BCryptPasswordEncoder con factor de costo por defecto"],
]
make_table(doc, ["Regla", "Detalle"], pol_rows, col_widths=[5, 12])

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = footer_p.add_run("GestionUsuario  ·  Documento generado automáticamente")
run_f.font.size = Pt(9)
run_f.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run_f.font.italic = True

# ── Guardar ───────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), "ModeloBaseDatos_GestionUsuario.docx")
doc.save(output_path)
print(f"✅ Documento generado exitosamente: {output_path}")
