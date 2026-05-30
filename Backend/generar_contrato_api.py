"""
Genera el documento Word con el Contrato de la API de GestionUsuario.
Uso: python3 generar_contrato_api.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Paleta de colores ─────────────────────────────────────────────────────────
AZUL_OSCURO   = RGBColor(0x1E, 0x3A, 0x5F)
AZUL_MEDIO    = RGBColor(0x2E, 0x6D, 0xA4)
AZUL_CLARO    = RGBColor(0xDB, 0xE9, 0xF5)
VERDE_OSCURO  = RGBColor(0x1A, 0x6B, 0x35)
VERDE_CLARO   = RGBColor(0xD4, 0xED, 0xDA)
NARANJA       = RGBColor(0xD4, 0x6B, 0x08)
NARANJA_CLARO = RGBColor(0xFD, 0xF0, 0xD5)
ROJO          = RGBColor(0xA8, 0x1C, 0x1C)
ROJO_CLARO    = RGBColor(0xF8, 0xD7, 0xDA)
MORADO        = RGBColor(0x5A, 0x2D, 0x82)
MORADO_CLARO  = RGBColor(0xE8, 0xD5, 0xF5)
GRIS_CLARO    = RGBColor(0xF8, 0xF9, 0xFA)
GRIS_FILA     = RGBColor(0xF0, 0xF4, 0xF8)
BLANCO        = RGBColor(0xFF, 0xFF, 0xFF)
NEGRO         = RGBColor(0x21, 0x25, 0x29)

# Colores por método HTTP
METHOD_COLORS = {
    "GET":    (RGBColor(0x19, 0x76, 0xD2), RGBColor(0xE3, 0xF2, 0xFD)),
    "POST":   (RGBColor(0x1B, 0x7A, 0x34), RGBColor(0xE8, 0xF5, 0xE9)),
    "PUT":    (RGBColor(0xE6, 0x5C, 0x00), RGBColor(0xFF, 0xF8, 0xE1)),
    "PATCH":  (RGBColor(0x6A, 0x1B, 0x9A), RGBColor(0xF3, 0xE5, 0xF5)),
    "DELETE": (RGBColor(0xC6, 0x28, 0x28), RGBColor(0xFF, 0xEB, 0xEE)),
}

# ── Helpers XML ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    tcPr.append(shd)

def add_borders(table, color="CCCCCC"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcB = OxmlElement("w:tcBorders")
            for side in ("top","left","bottom","right","insideH","insideV"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), color)
                tcB.append(b)
            tcPr.append(tcB)

def cell_para(cell, text, bold=False, italic=False, size=9,
              color=None, align=WD_ALIGN_PARAGRAPH.LEFT, font="Calibri"):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

def add_code_block(doc, code: str, lang_label="JSON"):
    """Añade un bloque de código con fondo gris y fuente monoespaciada."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, RGBColor(0xF4, 0xF4, 0xF4))
    add_borders(table, "DDDDDD")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code.strip())
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x24, 0x29, 0x2E)
    table.columns[0].width = Cm(16)
    return table

def add_method_badge(doc, method, path, summary, auth_tag=""):
    """Añade encabezado de endpoint con badge de color por método."""
    fg, bg = METHOD_COLORS.get(method, (AZUL_MEDIO, AZUL_CLARO))
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_borders(table, "E0E0E0")

    # Columna badge método
    c0 = table.rows[0].cells[0]
    set_cell_bg(c0, fg)
    cell_para(c0, method, bold=True, size=10, color=BLANCO,
              align=WD_ALIGN_PARAGRAPH.CENTER, font="Courier New")
    c0.width = Cm(1.8)

    # Columna ruta
    c1 = table.rows[0].cells[1]
    set_cell_bg(c1, bg)
    cell_para(c1, f"  {path}", bold=True, size=10, color=NEGRO, font="Courier New")
    c1.width = Cm(9)

    # Columna auth tag
    c2 = table.rows[0].cells[2]
    set_cell_bg(c2, bg)
    cell_para(c2, auth_tag, bold=False, size=8.5, color=RGBColor(0x55,0x55,0x55),
              align=WD_ALIGN_PARAGRAPH.RIGHT)
    c2.width = Cm(5.2)

    doc.add_paragraph()
    # Summary
    p = doc.add_paragraph()
    r = p.add_run(f"  {summary}")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x33,0x33,0x33)

def make_params_table(doc, headers, rows, col_widths):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(hdr.cells[i], AZUL_OSCURO)
        cell_para(hdr.cells[i], h, bold=True, size=9, color=BLANCO,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, row in enumerate(rows):
        bg = GRIS_FILA if ri % 2 == 0 else BLANCO
        for ci, val in enumerate(row):
            set_cell_bg(table.rows[ri+1].cells[ci], bg)
            cell_para(table.rows[ri+1].cells[ci], val, bold=(ci==0), size=9,
                      font="Courier New" if ci < 2 else "Calibri")
    add_borders(table)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Cm(col_widths[i])
    return table

def section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = AZUL_OSCURO
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 80)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENTO
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Cm(2.5)
section.left_margin = section.right_margin = Cm(2.5)

# ── PORTADA ───────────────────────────────────────────────────────────────────
for _ in range(3): doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("CONTRATO DE LA API")
r.bold = True; r.font.size = Pt(28); r.font.color.rgb = AZUL_OSCURO

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run("GestionUsuario — Auth & User Management")
r2.font.size = Pt(15); r2.font.color.rgb = AZUL_MEDIO

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = info.add_run("Spring Boot 3.3  ·  REST  ·  JWT Bearer  ·  JSON")
r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(0x66,0x66,0x66); r3.font.italic = True

doc.add_paragraph()
base = doc.add_paragraph()
base.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = base.add_run("Base URL (dev):  http://localhost:8091")
r4.bold = True; r4.font.size = Pt(12); r4.font.name = "Courier New"

doc.add_page_break()

# ── 1. CONVENCIONES GENERALES ─────────────────────────────────────────────────
doc.add_heading("1. Convenciones Generales", level=1)

conv_text = (
    "Todos los endpoints devuelven y consumen JSON (Content-Type: application/json). "
    "Las peticiones autenticadas deben incluir el header Authorization con el token JWT "
    "obtenido en el login.\n\n"
    "Authorization: Bearer <token>"
)
p = doc.add_paragraph(conv_text)
p.runs[0].font.size = Pt(10)

doc.add_paragraph()
doc.add_heading("1.1  Estructura de Error Estándar", level=2)
doc.add_paragraph("Todos los errores retornan el mismo esquema JSON:").runs[0].font.size = Pt(10)
add_code_block(doc, """{
  "timestamp": "2026-05-28T10:30:00.123",
  "status":    401,
  "error":     "Unauthorized",
  "message":   "Debes iniciar sesión para acceder a este recurso"
}""")

doc.add_paragraph()
doc.add_heading("1.2  Códigos de Estado HTTP", level=2)
http_rows = [
    ["200", "OK",                   "Petición exitosa"],
    ["201", "Created",              "Recurso creado exitosamente"],
    ["400", "Bad Request",          "Datos de entrada inválidos o incompletos"],
    ["401", "Unauthorized",         "Token ausente, inválido o expirado"],
    ["403", "Forbidden",            "Sin permisos para esta acción (cuenta inactiva o rol insuficiente)"],
    ["404", "Not Found",            "Recurso no encontrado"],
    ["409", "Conflict",             "Conflicto: correo o RUT ya registrado"],
    ["500", "Internal Server Error","Error inesperado del servidor"],
]
make_params_table(doc, ["Código","Estado","Descripción"], http_rows, [2, 4, 10])

doc.add_paragraph()
doc.add_heading("1.3  Niveles de Acceso", level=2)
access_rows = [
    ["🌐 Público",          "Sin token requerido",                       "/api/auth/login, /api/auth/register, etc."],
    ["🔐 Autenticado",      "Token JWT válido de cualquier rol",          "/api/auth/me, /api/auth/me/password"],
    ["👑 ROLE_ADMIN",       "Token JWT con rol ROLE_ADMIN",               "/api/admin/** (excepto /nombres)"],
    ["👁️  ADMIN/SUPERVISOR", "Token JWT con ROLE_ADMIN o ROLE_SUPERVISOR", "/api/admin/users/nombres"],
]
make_params_table(doc, ["Nivel","Requerimiento","Aplica en"], access_rows, [3.5, 5, 7.5])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. ENDPOINTS DE AUTENTICACIÓN
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Autenticación  —  /api/auth", level=1)
p = doc.add_paragraph("Endpoints públicos. No requieren token JWT salvo los indicados.")
p.runs[0].font.size = Pt(10)
doc.add_paragraph()

# ── 2.1 POST /api/auth/register ───────────────────────────────────────────────
doc.add_heading("2.1  Registrar Usuario", level=2)
add_method_badge(doc, "POST", "/api/auth/register",
                 "Crea una nueva cuenta de usuario. Queda inactiva hasta activación por admin.",
                 "🌐 Público")
doc.add_paragraph()

section_title(doc, "Request Body")
add_code_block(doc, """{
  "nombreCompleto": "Juan Pérez González",   // string, requerido
  "email":          "juan@example.com",      // string email, requerido, único
  "password":       "MiPass123!",            // string, requerido — ver política §5
  "rut":            "12345678-9",            // string, opcional, único
  "telefono":       "+56912345678",          // string, opcional
  "roles":          ["ROLE_USER"]            // array, opcional — default: ROLE_USER
}""")

doc.add_paragraph()
section_title(doc, "Response  201 Created")
add_code_block(doc, """{
  "message": "Usuario registrado exitosamente. Espera la activación de tu cuenta."
}""")

doc.add_paragraph()
section_title(doc, "Errores posibles")
err_rows = [
    ["400", "Campos requeridos vacíos o email con formato inválido"],
    ["400", "La contraseña no cumple la política de complejidad"],
    ["409", "El correo ingresado ya se encuentra registrado"],
    ["409", "El RUT ingresado ya se encuentra registrado"],
]
make_params_table(doc, ["Código","Mensaje"], err_rows, [2, 14])
divider(doc)

# ── 2.2 POST /api/auth/login ──────────────────────────────────────────────────
doc.add_heading("2.2  Iniciar Sesión (Login)", level=2)
add_method_badge(doc, "POST", "/api/auth/login",
                 "Autentica con email y contraseña. Retorna JWT + permisos del usuario.",
                 "🌐 Público")
doc.add_paragraph()

section_title(doc, "Request Body")
add_code_block(doc, """{
  "email":    "admin@gestion.com",   // string email, requerido
  "password": "AdminPass123!"        // string, requerido
}""")

doc.add_paragraph()
section_title(doc, "Response  200 OK")
add_code_block(doc, """{
  "token":          "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...",
  "type":           "Bearer",
  "email":          "admin@gestion.com",
  "nombreCompleto": "Administrador Maestro",
  "rol":            "ROLE_ADMIN",
  "permisos": [
    {
      "modulo":        "Gestión de Usuarios",
      "rutaFrontend":  "/usuarios",
      "acciones":      ["CREAR", "LEER", "ACTUALIZAR", "ELIMINAR"]
    },
    {
      "modulo":        "Gestión de Roles",
      "rutaFrontend":  "/roles",
      "acciones":      ["CREAR", "LEER", "ACTUALIZAR", "ELIMINAR"]
    }
  ]
}""")

doc.add_paragraph()
section_title(doc, "Errores posibles")
err_rows2 = [
    ["400", "Cuerpo de la petición vacío o mal formado"],
    ["401", "El correo o la contraseña son incorrectos"],
    ["403", "Tu cuenta está registrada pero debe ser activada por un administrador"],
]
make_params_table(doc, ["Código","Mensaje"], err_rows2, [2, 14])
divider(doc)

# ── 2.3 POST /api/auth/logout ─────────────────────────────────────────────────
doc.add_heading("2.3  Cerrar Sesión (Logout)", level=2)
add_method_badge(doc, "POST", "/api/auth/logout",
                 "Limpia el contexto de seguridad del servidor. El cliente debe eliminar el token.",
                 "🌐 Público")
doc.add_paragraph()

section_title(doc, "Request Body")
doc.add_paragraph("  Sin cuerpo requerido.").runs[0].font.size = Pt(9)

doc.add_paragraph()
section_title(doc, "Response  200 OK")
add_code_block(doc, '{\n  "message": "Sesión cerrada correctamente"\n}')
divider(doc)

doc.add_page_break()

# ── 2.4 POST /api/auth/reset-password ────────────────────────────────────────
doc.add_heading("2.4  Solicitar Recuperación de Contraseña", level=2)
add_method_badge(doc, "POST", "/api/auth/reset-password",
                 "Genera un token UUID (TTL 15 min) y lo imprime en consola (simulación de email).",
                 "🌐 Público")
doc.add_paragraph()
p_nota = doc.add_paragraph(
    "  ⚠  Por seguridad, siempre retorna 200 OK aunque el correo no exista (evita enumeración de usuarios).")
p_nota.runs[0].font.size = Pt(9); p_nota.runs[0].font.italic = True
doc.add_paragraph()

section_title(doc, "Request Body")
add_code_block(doc, '{\n  "email": "juan@example.com"   // string email, requerido\n}')

doc.add_paragraph()
section_title(doc, "Response  200 OK")
add_code_block(doc, '{\n  "message": "Si el correo existe, se ha enviado un enlace de recuperación."\n}')
divider(doc)

# ── 2.5 POST /api/auth/new-password ──────────────────────────────────────────
doc.add_heading("2.5  Confirmar Nueva Contraseña", level=2)
add_method_badge(doc, "POST", "/api/auth/new-password",
                 "Valida el token de reset y establece la nueva contraseña. El token queda destruido.",
                 "🌐 Público")
doc.add_paragraph()

section_title(doc, "Request Body")
add_code_block(doc, """{
  "token":       "3f7a2b1c-8e4d-4a9f-b2c1-...",  // string UUID, requerido
  "newPassword": "NuevaClave456!"                 // string, requerido — ver política §5
}""")

doc.add_paragraph()
section_title(doc, "Response  200 OK")
add_code_block(doc, '{\n  "message": "Contraseña actualizada exitosamente"\n}')

doc.add_paragraph()
section_title(doc, "Errores posibles")
err_rows5 = [
    ["400", "Token inválido o no encontrado"],
    ["400", "El token ha expirado. Solicita uno nuevo"],
    ["400", "La contraseña no cumple la política de complejidad"],
]
make_params_table(doc, ["Código","Mensaje"], err_rows5, [2, 14])
divider(doc)

doc.add_page_break()

# ── 2.6 GET /api/auth/me ─────────────────────────────────────────────────────
doc.add_heading("2.6  Ver Perfil Propio", level=2)
add_method_badge(doc, "GET", "/api/auth/me",
                 "Retorna el perfil completo del usuario autenticado.",
                 "🔐 Autenticado")
doc.add_paragraph()

section_title(doc, "Response  200 OK")
add_code_block(doc, """{
  "id":             1,
  "nombreCompleto": "Administrador Maestro",
  "email":          "admin@gestion.com",
  "rut":            null,
  "telefono":       null,
  "rol":            "ROLE_ADMIN",
  "activo":         true,
  "fechaCreacion":  "2026-05-28T10:00:00"
}""")

doc.add_paragraph()
section_title(doc, "Errores posibles")
make_params_table(doc, ["Código","Mensaje"],
    [["401","Debes iniciar sesión para acceder a este recurso"]], [2, 14])
divider(doc)

# ── 2.7 PUT /api/auth/me ─────────────────────────────────────────────────────
doc.add_heading("2.7  Actualizar Perfil Propio", level=2)
add_method_badge(doc, "PUT", "/api/auth/me",
                 "Actualiza nombreCompleto y/o teléfono del usuario autenticado.",
                 "🔐 Autenticado")
doc.add_paragraph()

section_title(doc, "Request Body  (al menos un campo)")
add_code_block(doc, """{
  "nombreCompleto": "Juan Pérez Actualizado",   // string, opcional
  "telefono":       "+56987654321"              // string, opcional
}""")

doc.add_paragraph()
section_title(doc, "Response  200 OK")
add_code_block(doc, '{\n  "message": "Perfil actualizado correctamente"\n}')

doc.add_paragraph()
section_title(doc, "Errores posibles")
make_params_table(doc, ["Código","Mensaje"], [
    ["400","Debes enviar al menos un campo para actualizar: nombreCompleto o telefono"],
    ["401","Debes iniciar sesión para acceder a este recurso"],
], [2, 14])
divider(doc)

# ── 2.8 PUT /api/auth/me/password ─────────────────────────────────────────────
doc.add_heading("2.8  Cambiar Contraseña Propia", level=2)
add_method_badge(doc, "PUT", "/api/auth/me/password",
                 "Cambia la contraseña validando la actual. Aplica política de complejidad.",
                 "🔐 Autenticado")
doc.add_paragraph()

section_title(doc, "Request Body")
add_code_block(doc, """{
  "currentPassword": "AdminPass123!",   // string, requerido
  "newPassword":     "NuevoClave789!"   // string, requerido — ver política §5
}""")

doc.add_paragraph()
section_title(doc, "Response  200 OK")
add_code_block(doc, '{\n  "message": "Contraseña actualizada correctamente"\n}')

doc.add_paragraph()
section_title(doc, "Errores posibles")
make_params_table(doc, ["Código","Mensaje"], [
    ["400","Debes ingresar tu contraseña actual"],
    ["400","Debes ingresar la nueva contraseña"],
    ["400","La contraseña actual es incorrecta"],
    ["400","La nueva contraseña no puede ser igual a la actual"],
    ["400","La contraseña no cumple la política de complejidad"],
    ["401","Debes iniciar sesión para acceder a este recurso"],
], [2, 14])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. ENDPOINTS DE ADMINISTRACIÓN
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Administración  —  /api/admin", level=1)
p = doc.add_paragraph(
    "Todos los endpoints de este grupo requieren ROLE_ADMIN excepto donde se indica. "
    "Header requerido: Authorization: Bearer <token>")
p.runs[0].font.size = Pt(10)
doc.add_paragraph()

# ── 3.1 GET /api/admin/users ──────────────────────────────────────────────────
doc.add_heading("3.1  Listar Todos los Usuarios", level=2)
add_method_badge(doc, "GET", "/api/admin/users",
                 "Retorna la lista completa de usuarios con todos sus campos.", "👑 ROLE_ADMIN")
doc.add_paragraph()

section_title(doc, "Response  200 OK")
add_code_block(doc, """[
  {
    "id":             1,
    "role":           { "id": 1, "nombre": "ROLE_ADMIN", "descripcion": "Administrador del sistema" },
    "nombreCompleto": "Administrador Maestro",
    "email":          "admin@gestion.com",
    "rut":            null,
    "telefono":       null,
    "isActive":       true,
    "createdAt":      "2026-05-28T10:00:00"
  },
  {
    "id":             2,
    "role":           { "id": 2, "nombre": "ROLE_USER", "descripcion": "Usuario estándar" },
    "nombreCompleto": "Usuario Estándar",
    "email":          "user@gestion.com",
    "rut":            null,
    "telefono":       null,
    "isActive":       true,
    "createdAt":      "2026-05-28T10:00:00"
  }
]""")
divider(doc)

# ── 3.2 GET /api/admin/users/nombres ─────────────────────────────────────────
doc.add_heading("3.2  Mapa ID → Nombre de Usuarios", level=2)
add_method_badge(doc, "GET", "/api/admin/users/nombres",
                 "Retorna id y nombre de todos los usuarios. Útil para resolver IDs.",
                 "👁️  ADMIN / SUPERVISOR")
doc.add_paragraph()

section_title(doc, "Response  200 OK")
add_code_block(doc, """[
  { "id": 1, "nombre": "Administrador Maestro" },
  { "id": 2, "nombre": "Usuario Estándar" },
  { "id": 3, "nombre": "Supervisor del Sistema" }
]""")
divider(doc)

# ── 3.3 POST /api/admin/users ────────────────────────────────────────────────
doc.add_heading("3.3  Crear Usuario (Admin)", level=2)
add_method_badge(doc, "POST", "/api/admin/users",
                 "Crea un usuario directamente activo (aprovisionamiento directo).", "👑 ROLE_ADMIN")
doc.add_paragraph()

section_title(doc, "Request Body")
add_code_block(doc, """{
  "nombreCompleto": "María García López",    // string, requerido
  "email":          "maria@example.com",    // string email, requerido, único
  "password":       "MiPass123!",           // string, requerido
  "rut":            "98765432-1",           // string, opcional, único
  "telefono":       "+56911111111",         // string, opcional
  "roles":          ["ROLE_SUPERVISOR"]     // array, opcional — default: ROLE_USER
}""")

doc.add_paragraph()
section_title(doc, "Response  201 Created  — objeto User completo")
add_code_block(doc, """{
  "id":             4,
  "role":           { "id": 3, "nombre": "ROLE_SUPERVISOR", "descripcion": "..." },
  "nombreCompleto": "María García López",
  "email":          "maria@example.com",
  "password":       "$2a$10$...",
  "rut":            "98765432-1",
  "telefono":       "+56911111111",
  "isActive":       true,
  "createdAt":      "2026-05-28T12:00:00"
}""")

doc.add_paragraph()
section_title(doc, "Errores posibles")
make_params_table(doc, ["Código","Mensaje"], [
    ["400","Campos requeridos vacíos o email inválido"],
    ["400","La contraseña no cumple la política de complejidad"],
    ["409","El correo ingresado ya se encuentra registrado"],
    ["409","El RUT ingresado ya se encuentra registrado"],
], [2, 14])
divider(doc)

doc.add_page_break()

# ── 3.4 PUT /api/admin/users/{id} ────────────────────────────────────────────
doc.add_heading("3.4  Editar Usuario", level=2)
add_method_badge(doc, "PUT", "/api/admin/users/{id}",
                 "Actualiza los campos del usuario. Todos los campos son opcionales, pero se requiere al menos uno.",
                 "👑 ROLE_ADMIN")
doc.add_paragraph()

section_title(doc, "Path Parameter")
make_params_table(doc, ["Parámetro","Tipo","Descripción"],
    [["id","Long","ID único del usuario a modificar"]], [3.5, 2.5, 10])

doc.add_paragraph()
section_title(doc, "Request Body  (al menos un campo)")
add_code_block(doc, """{
  "email":          "nuevo@example.com",   // string email, opcional
  "nombreCompleto": "Nombre Actualizado",  // string, opcional
  "rut":            "11111111-1",          // string, opcional
  "telefono":       "+56922222222",        // string, opcional
  "password":       "NuevaClave789!"       // string, opcional — aplica política
}""")

doc.add_paragraph()
section_title(doc, "Response  200 OK")
add_code_block(doc, '{\n  "message": "Usuario actualizado correctamente"\n}')

doc.add_paragraph()
section_title(doc, "Errores posibles")
make_params_table(doc, ["Código","Mensaje"], [
    ["400","Debes enviar al menos un campo para actualizar"],
    ["404","Usuario no encontrado"],
    ["409","El nuevo correo ya está en uso"],
    ["409","El RUT ya está registrado"],
], [2, 14])
divider(doc)

# ── 3.5 PATCH /api/admin/users/{id}/toggle-status ────────────────────────────
doc.add_heading("3.5  Activar / Desactivar Usuario", level=2)
add_method_badge(doc, "PATCH", "/api/admin/users/{id}/toggle-status",
                 "Invierte el estado activo del usuario (soft delete / reactivación).",
                 "👑 ROLE_ADMIN")
doc.add_paragraph()

section_title(doc, "Path Parameter")
make_params_table(doc, ["Parámetro","Tipo","Descripción"],
    [["id","Long","ID del usuario a activar o desactivar"]], [3.5, 2.5, 10])

doc.add_paragraph()
section_title(doc, "Request Body")
doc.add_paragraph("  Sin cuerpo requerido.").runs[0].font.size = Pt(9)

doc.add_paragraph()
section_title(doc, "Response  200 OK")
add_code_block(doc, '{\n  "message": "Estado del usuario actualizado"\n}')

doc.add_paragraph()
section_title(doc, "Errores posibles")
make_params_table(doc, ["Código","Mensaje"],
    [["404","Usuario no encontrado"]], [2, 14])
divider(doc)

doc.add_page_break()

# ── 3.6 POST /api/admin/users/{userId}/roles/{roleName} ──────────────────────
doc.add_heading("3.6  Asignar Rol a Usuario", level=2)
add_method_badge(doc, "POST", "/api/admin/users/{userId}/roles/{roleName}",
                 "Asigna un rol al usuario especificado.", "👑 ROLE_ADMIN")
doc.add_paragraph()

section_title(doc, "Path Parameters")
make_params_table(doc, ["Parámetro","Tipo","Descripción"], [
    ["userId",   "Long",   "ID del usuario"],
    ["roleName", "String", "Nombre exacto del rol (ej. ROLE_SUPERVISOR)"],
], [3.5, 2.5, 10])

doc.add_paragraph()
section_title(doc, "Response  200 OK")
add_code_block(doc, '{\n  "message": "Rol asignado correctamente"\n}')

doc.add_paragraph()
section_title(doc, "Errores posibles")
make_params_table(doc, ["Código","Mensaje"], [
    ["404","Usuario no encontrado"],
    ["404","Rol no encontrado: ROLE_XXX"],
], [2, 14])
divider(doc)

# ── 3.7 DELETE /api/admin/users/{userId}/roles/{roleName} ────────────────────
doc.add_heading("3.7  Revocar Rol a Usuario", level=2)
add_method_badge(doc, "DELETE", "/api/admin/users/{userId}/roles/{roleName}",
                 "Degrada al usuario a ROLE_USER (rol base del sistema).", "👑 ROLE_ADMIN")
doc.add_paragraph()

section_title(doc, "Path Parameters")
make_params_table(doc, ["Parámetro","Tipo","Descripción"], [
    ["userId",   "Long",   "ID del usuario"],
    ["roleName", "String", "Rol a revocar (ej. ROLE_ADMIN)"],
], [3.5, 2.5, 10])

doc.add_paragraph()
section_title(doc, "Response  200 OK")
add_code_block(doc, '{\n  "message": "Rol ROLE_ADMIN revocado al usuario 2"\n}')
divider(doc)

doc.add_page_break()

# ── 3.8 GET /api/admin/roles ──────────────────────────────────────────────────
doc.add_heading("3.8  Listar Roles", level=2)
add_method_badge(doc, "GET", "/api/admin/roles",
                 "Retorna todos los roles del sistema con sus permisos.", "👑 ROLE_ADMIN")
doc.add_paragraph()

section_title(doc, "Response  200 OK")
add_code_block(doc, """[
  {
    "id":          1,
    "nombre":      "ROLE_ADMIN",
    "descripcion": "Administrador del sistema",
    "permisos": [
      {
        "modulo":  { "id": 1, "nombre": "Gestión de Usuarios", "rutaFrontend": "/usuarios" },
        "acceso":  { "id": 1, "nombre": "CREAR" }
      },
      {
        "modulo":  { "id": 1, "nombre": "Gestión de Usuarios", "rutaFrontend": "/usuarios" },
        "acceso":  { "id": 2, "nombre": "LEER" }
      }
    ]
  },
  {
    "id":          2,
    "nombre":      "ROLE_USER",
    "descripcion": "Usuario estándar",
    "permisos":    []
  }
]""")
divider(doc)

# ── 3.9 POST /api/admin/roles ─────────────────────────────────────────────────
doc.add_heading("3.9  Crear Rol", level=2)
add_method_badge(doc, "POST", "/api/admin/roles",
                 "Crea un nuevo rol. Se agrega el prefijo ROLE_ automáticamente si no está presente.",
                 "👑 ROLE_ADMIN")
doc.add_paragraph()

section_title(doc, "Query Parameter")
make_params_table(doc, ["Parámetro","Tipo","Descripción"],
    [["roleName","String","Nombre del rol (ej. SUPERVISOR → se guarda como ROLE_SUPERVISOR)"]], [3.5, 2.5, 10])

doc.add_paragraph()
section_title(doc, "Response  200 OK")
add_code_block(doc, '{\n  "message": "Rol creado exitosamente"\n}')

doc.add_paragraph()
section_title(doc, "Errores posibles")
make_params_table(doc, ["Código","Mensaje"], [
    ["400","Debes ingresar el nombre del rol"],
    ["400","El rol 'ROLE_XXX' ya existe"],
], [2, 14])
divider(doc)

doc.add_page_break()

# ── 3.10 GET /api/admin/modules ───────────────────────────────────────────────
doc.add_heading("3.10  Listar Módulos", level=2)
add_method_badge(doc, "GET", "/api/admin/modules",
                 "Retorna todos los módulos del sistema.", "👑 ROLE_ADMIN")
doc.add_paragraph()

section_title(doc, "Response  200 OK")
add_code_block(doc, """[
  { "id": 1, "nombre": "Gestión de Usuarios", "descripcion": "...", "rutaFrontend": "/usuarios" },
  { "id": 2, "nombre": "Gestión de Roles",    "descripcion": "...", "rutaFrontend": "/roles"    }
]""")
divider(doc)

# ── 3.11 GET /api/admin/access-types ─────────────────────────────────────────
doc.add_heading("3.11  Listar Tipos de Acceso", level=2)
add_method_badge(doc, "GET", "/api/admin/access-types",
                 "Retorna todos los tipos de acceso disponibles.", "👑 ROLE_ADMIN")
doc.add_paragraph()

section_title(doc, "Response  200 OK")
add_code_block(doc, """[
  { "id": 1, "nombre": "CREAR",      "descripcion": "Permite crear registros"              },
  { "id": 2, "nombre": "LEER",       "descripcion": "Permite visualizar información"       },
  { "id": 3, "nombre": "ACTUALIZAR", "descripcion": "Permite modificar registros"          },
  { "id": 4, "nombre": "ELIMINAR",   "descripcion": "Permite eliminar o desactivar registros" }
]""")
divider(doc)

# ── 3.12 POST /api/admin/roles/{idRol}/permissions ────────────────────────────
doc.add_heading("3.12  Otorgar Permiso a Rol", level=2)
add_method_badge(doc, "POST", "/api/admin/roles/{idRol}/permissions",
                 "Agrega un permiso (módulo + acción) al rol indicado.", "👑 ROLE_ADMIN")
doc.add_paragraph()

section_title(doc, "Path Parameter")
make_params_table(doc, ["Parámetro","Tipo","Descripción"],
    [["idRol","Long","ID del rol al que se otorgará el permiso"]], [3.5, 2.5, 10])

doc.add_paragraph()
section_title(doc, "Request Body")
add_code_block(doc, """{
  "idModulo": 1,   // Long — ID del módulo (ver /api/admin/modules)
  "idAcceso": 2    // Long — ID del tipo de acceso (ver /api/admin/access-types)
}""")

doc.add_paragraph()
section_title(doc, "Response  200 OK")
add_code_block(doc, '{\n  "message": "Permiso otorgado exitosamente al rol"\n}')

doc.add_paragraph()
section_title(doc, "Errores posibles")
make_params_table(doc, ["Código","Mensaje"], [
    ["404","Rol no encontrado"],
    ["404","Módulo no encontrado"],
    ["404","Acceso no encontrado"],
], [2, 14])
divider(doc)

# ── 3.13 DELETE /api/admin/roles/{idRol}/permissions ─────────────────────────
doc.add_heading("3.13  Revocar Permiso de Rol", level=2)
add_method_badge(doc, "DELETE", "/api/admin/roles/{idRol}/permissions",
                 "Elimina un permiso (módulo + acción) del rol indicado.", "👑 ROLE_ADMIN")
doc.add_paragraph()

section_title(doc, "Path Parameter")
make_params_table(doc, ["Parámetro","Tipo","Descripción"],
    [["idRol","Long","ID del rol al que se revocará el permiso"]], [3.5, 2.5, 10])

doc.add_paragraph()
section_title(doc, "Request Body")
add_code_block(doc, """{
  "idModulo": 1,   // Long — ID del módulo
  "idAcceso": 2    // Long — ID del tipo de acceso
}""")

doc.add_paragraph()
section_title(doc, "Response  200 OK")
add_code_block(doc, '{\n  "message": "Permiso revocado exitosamente del rol"\n}')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. RESUMEN DE ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Resumen Completo de Endpoints", level=1)

summary_rows = [
    # Auth
    ["POST",   "/api/auth/register",                         "Registrar usuario",               "🌐 Público"],
    ["POST",   "/api/auth/login",                            "Iniciar sesión → JWT",            "🌐 Público"],
    ["POST",   "/api/auth/logout",                           "Cerrar sesión",                   "🌐 Público"],
    ["POST",   "/api/auth/reset-password",                   "Solicitar reset de contraseña",   "🌐 Público"],
    ["POST",   "/api/auth/new-password",                     "Confirmar nueva contraseña",      "🌐 Público"],
    ["GET",    "/api/auth/me",                               "Ver perfil propio",               "🔐 Autenticado"],
    ["PUT",    "/api/auth/me",                               "Actualizar perfil propio",        "🔐 Autenticado"],
    ["PUT",    "/api/auth/me/password",                      "Cambiar contraseña propia",       "🔐 Autenticado"],
    # Admin - Usuarios
    ["GET",    "/api/admin/users",                           "Listar todos los usuarios",       "👑 ROLE_ADMIN"],
    ["GET",    "/api/admin/users/nombres",                   "Mapa id→nombre",                  "👁️  ADMIN/SUPERVISOR"],
    ["POST",   "/api/admin/users",                           "Crear usuario",                   "👑 ROLE_ADMIN"],
    ["PUT",    "/api/admin/users/{id}",                      "Editar usuario",                  "👑 ROLE_ADMIN"],
    ["PATCH",  "/api/admin/users/{id}/toggle-status",        "Activar / Desactivar",            "👑 ROLE_ADMIN"],
    ["POST",   "/api/admin/users/{userId}/roles/{roleName}", "Asignar rol",                     "👑 ROLE_ADMIN"],
    ["DELETE", "/api/admin/users/{userId}/roles/{roleName}", "Revocar rol",                     "👑 ROLE_ADMIN"],
    # Admin - Roles y permisos
    ["GET",    "/api/admin/roles",                           "Listar roles",                    "👑 ROLE_ADMIN"],
    ["POST",   "/api/admin/roles",                           "Crear rol",                       "👑 ROLE_ADMIN"],
    ["GET",    "/api/admin/modules",                         "Listar módulos",                  "👑 ROLE_ADMIN"],
    ["GET",    "/api/admin/access-types",                    "Listar tipos de acceso",          "👑 ROLE_ADMIN"],
    ["POST",   "/api/admin/roles/{idRol}/permissions",       "Otorgar permiso a rol",           "👑 ROLE_ADMIN"],
    ["DELETE", "/api/admin/roles/{idRol}/permissions",       "Revocar permiso de rol",          "👑 ROLE_ADMIN"],
]

table = doc.add_table(rows=1+len(summary_rows), cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0]
for i, h in enumerate(["Método", "Endpoint", "Descripción", "Acceso"]):
    set_cell_bg(hdr.cells[i], AZUL_OSCURO)
    cell_para(hdr.cells[i], h, bold=True, size=9, color=BLANCO, align=WD_ALIGN_PARAGRAPH.CENTER)

for ri, row_data in enumerate(summary_rows):
    method = row_data[0]
    fg, bg = METHOD_COLORS.get(method, (AZUL_MEDIO, AZUL_CLARO))
    tr = table.rows[ri + 1]
    # Método
    set_cell_bg(tr.cells[0], fg)
    cell_para(tr.cells[0], method, bold=True, size=8.5, color=BLANCO,
              align=WD_ALIGN_PARAGRAPH.CENTER, font="Courier New")
    # Endpoint
    row_bg = GRIS_FILA if ri % 2 == 0 else BLANCO
    set_cell_bg(tr.cells[1], row_bg)
    cell_para(tr.cells[1], row_data[1], size=8.5, font="Courier New")
    # Descripción
    set_cell_bg(tr.cells[2], row_bg)
    cell_para(tr.cells[2], row_data[2], size=9)
    # Acceso
    set_cell_bg(tr.cells[3], row_bg)
    cell_para(tr.cells[3], row_data[3], size=8.5)

add_borders(table)
col_w = [1.8, 6.5, 5.5, 3.2]
for row in table.rows:
    for i, cell in enumerate(row.cells):
        cell.width = Cm(col_w[i])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. POLÍTICA DE CONTRASEÑAS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Política de Contraseñas", level=1)
doc.add_paragraph(
    "Aplica en: POST /register, POST /new-password, PUT /me/password, "
    "POST /admin/users, PUT /admin/users/{id}"
).runs[0].font.size = Pt(10)
doc.add_paragraph()

make_params_table(doc, ["Regla","Detalle","Ejemplo válido"], [
    ["Longitud mínima",   "8 caracteres como mínimo",              "MiPass1!"],
    ["Mayúscula",         "Al menos 1 letra mayúscula (A-Z)",      "mipass1! ❌ → Mipass1! ✅"],
    ["Minúscula",         "Al menos 1 letra minúscula (a-z)",      "MIPASS1! ❌ → MiPass1! ✅"],
    ["Número",            "Al menos 1 dígito (0-9)",               "MiPass!! ❌ → MiPass1! ✅"],
    ["Carácter especial", "Al menos 1 de: !@#$%^&*()_+-=[]{};etc","MiPass12 ❌ → MiPass1! ✅"],
], [4, 6, 6])

doc.add_paragraph()
section_title(doc, "Respuesta de error cuando no se cumple la política (400 Bad Request)")
add_code_block(doc, """{
  "timestamp": "2026-05-28T10:00:00",
  "status":    400,
  "error":     "Bad Request",
  "message":   "La contraseña debe contener al menos un carácter especial (!@#$%^&* etc.)"
}""")

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = fp.add_run("GestionUsuario  ·  Contrato de API  ·  Documento generado automáticamente")
rf.font.size = Pt(8.5); rf.font.italic = True
rf.font.color.rgb = RGBColor(0x99,0x99,0x99)

# ── Guardar ───────────────────────────────────────────────────────────────────
output = os.path.join(os.path.dirname(__file__), "ContratoAPI_GestionUsuario.docx")
doc.save(output)
print(f"✅ Contrato de API generado: {output}")
